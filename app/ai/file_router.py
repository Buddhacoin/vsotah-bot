import asyncio
import base64
import os
import csv
import json
from io import BytesIO, StringIO
from typing import Any

from openai import AsyncOpenAI

from app.ai.prompts import clean_ai_answer, system_prompt
from app.ai.personalities import get_personality
from app.ai.router import ai_router
from app.ai.vision_router import build_pdf_vision_prompt


OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_VISION_MODEL = os.getenv("OPENAI_VISION_MODEL", "gpt-5.4-mini")
FILE_TEXT_LIMIT = int(os.getenv("FILE_TEXT_LIMIT", "22000"))
FILE_MAX_PAGES = int(os.getenv("FILE_MAX_PAGES", "35"))
FILE_XLSX_MAX_ROWS = int(os.getenv("FILE_XLSX_MAX_ROWS", "120"))
FILE_XLSX_MAX_SHEETS = int(os.getenv("FILE_XLSX_MAX_SHEETS", "6"))
PDF_VISION_MAX_PAGES = int(os.getenv("PDF_VISION_MAX_PAGES", "4"))
TEXT_HISTORY_LIMIT = int(os.getenv("TEXT_HISTORY_LIMIT", "6"))
VISION_HISTORY_LIMIT = int(os.getenv("VISION_HISTORY_LIMIT", "2"))
VISION_MAX_TOKENS = int(os.getenv("VISION_MAX_TOKENS", "900"))
AI_TIMEOUT_SECONDS = int(os.getenv("AI_TIMEOUT_SECONDS", "75"))

openai_client = AsyncOpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None


SUPPORTED_FILE_TYPES = {
    "txt", "md", "csv", "log", "json", "html", "htm",
    "pdf", "docx", "xlsx", "xlsm",
}


def _extension(filename: str) -> str:
    return filename.lower().rsplit(".", 1)[-1] if "." in filename else ""


def _limit_text(text: str) -> str:
    return (text or "")[:FILE_TEXT_LIMIT]


def _decode_text_raw(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return file_bytes.decode(encoding, errors="replace")
        except Exception:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _decode_text_file(file_bytes: bytes) -> str:
    return _limit_text(_decode_text_raw(file_bytes))


def _extract_json_text(file_bytes: bytes) -> str:
    raw = _decode_text_raw(file_bytes).strip()
    try:
        data = json.loads(raw)
        pretty = json.dumps(data, ensure_ascii=False, indent=2)
        return _limit_text(pretty)
    except Exception:
        return _limit_text(raw)


def _extract_csv_text(file_bytes: bytes) -> str:
    raw = _decode_text_raw(file_bytes)
    sample = raw[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample)
    except Exception:
        dialect = csv.excel

    reader = csv.reader(StringIO(raw), dialect)
    lines: list[str] = []
    for index, row in enumerate(reader):
        if index == 0:
            lines.append("--- CSV headers / first row ---")
        if index >= FILE_XLSX_MAX_ROWS:
            lines.append(f"--- Показаны первые {FILE_XLSX_MAX_ROWS} строк CSV ---")
            break
        values = [str(cell).strip() for cell in row]
        if any(values):
            lines.append(" | ".join(values))
        if len("\n".join(lines)) >= FILE_TEXT_LIMIT:
            break
    return _limit_text("\n".join(lines) or raw)


def _extract_pdf_text(file_bytes: bytes) -> str:
    pages: list[str] = []

    # 1) pypdf — fast for regular PDFs with a text layer.
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(file_bytes))
        for index, page in enumerate(reader.pages[:FILE_MAX_PAGES], start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- Страница {index} ---\n{text.strip()}")
            if len("\n\n".join(pages)) >= FILE_TEXT_LIMIT:
                return _limit_text("\n\n".join(pages))
    except Exception as e:
        print(f"PDF PYPDF READ ERROR: {e}")

    # 2) pdfplumber — often better for tables and invoices.
    if not "\n".join(pages).strip():
        try:
            import pdfplumber
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for index, page in enumerate(pdf.pages[:FILE_MAX_PAGES], start=1):
                    text = page.extract_text() or ""
                    tables = []
                    try:
                        for table in page.extract_tables() or []:
                            rows = [" | ".join(str(cell or "") for cell in row) for row in table[:30]]
                            if rows:
                                tables.append("\n".join(rows))
                    except Exception:
                        pass
                    page_text = "\n".join(part for part in [text.strip(), *tables] if part.strip())
                    if page_text:
                        pages.append(f"--- Страница {index} ---\n{page_text}")
                    if len("\n\n".join(pages)) >= FILE_TEXT_LIMIT:
                        return _limit_text("\n\n".join(pages))
        except Exception as e:
            print(f"PDF PDFPLUMBER READ ERROR: {e}")

    # 3) PyMuPDF fallback.
    if not "\n".join(pages).strip():
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for index, page in enumerate(doc[:FILE_MAX_PAGES], start=1):
                text = page.get_text("text") or ""
                if text.strip():
                    pages.append(f"--- Страница {index} ---\n{text.strip()}")
                if len("\n\n".join(pages)) >= FILE_TEXT_LIMIT:
                    break
            doc.close()
        except Exception as e:
            print(f"PDF FITZ TEXT READ ERROR: {e}")

    return _limit_text("\n\n".join(pages))


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table_index, table in enumerate(doc.tables, start=1):
            parts.append(f"\n--- Таблица {table_index} ---")
            for row in table.rows[:80]:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
                if len("\n".join(parts)) >= FILE_TEXT_LIMIT:
                    return _limit_text("\n".join(parts))

        return _limit_text("\n".join(parts))
    except Exception as e:
        raise ValueError(f"DOCX_READ_ERROR: {e}")


def _extract_xlsx_text(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(BytesIO(file_bytes), data_only=True, read_only=True)
        lines: list[str] = []

        for ws in wb.worksheets[:FILE_XLSX_MAX_SHEETS]:
            lines.append(f"--- Лист: {ws.title} | строк: {ws.max_row} | колонок: {ws.max_column} ---")
            non_empty_rows = 0
            for row in ws.iter_rows(max_row=FILE_XLSX_MAX_ROWS, values_only=True):
                values = [str(v) if v is not None else "" for v in row]
                if any(v.strip() for v in values):
                    non_empty_rows += 1
                    lines.append(" | ".join(values))
                if len("\n".join(lines)) >= FILE_TEXT_LIMIT:
                    return _limit_text("\n".join(lines))
            lines.append(f"--- В выборке непустых строк: {non_empty_rows} ---")

        return _limit_text("\n".join(lines))
    except Exception as e:
        raise ValueError(f"XLSX_READ_ERROR: {e}")


def extract_document_text(filename: str, file_bytes: bytes) -> str:
    """Extracts readable text/data from user files for Files Pipeline 1.0."""
    ext = _extension(filename)

    if ext not in SUPPORTED_FILE_TYPES:
        raise ValueError("UNSUPPORTED_FILE_TYPE")

    if ext == "csv":
        return _extract_csv_text(file_bytes)

    if ext == "json":
        return _extract_json_text(file_bytes)

    if ext in {"txt", "md", "log", "html", "htm"}:
        return _decode_text_file(file_bytes)

    if ext == "pdf":
        return _extract_pdf_text(file_bytes)

    if ext == "docx":
        return _extract_docx_text(file_bytes)

    if ext in {"xlsx", "xlsm"}:
        return _extract_xlsx_text(file_bytes)

    raise ValueError("UNSUPPORTED_FILE_TYPE")


def detect_file_kind(filename: str, extracted_text: str = "") -> str:
    ext = _extension(filename)
    lower = f"{filename}\n{extracted_text[:2000]}".lower()

    if ext in {"xlsx", "xlsm", "csv"}:
        return "spreadsheet"
    if ext == "pdf" and not extracted_text.strip():
        return "scanned_pdf"
    if any(word in lower for word in ("договор", "contract", "agreement", "акт", "счет", "invoice")):
        return "business_document"
    if any(word in lower for word in ("диплом", "курсов", "реферат", "dissertation", "thesis")):
        return "study_document"
    if ext == "docx":
        return "word_document"
    if ext == "pdf":
        return "pdf_document"
    return "text_document"


def detect_file_task(question: str, filename: str, extracted_text: str = "") -> str:
    lower = f"{question}\n{filename}\n{extracted_text[:1200]}".lower()
    if any(x in lower for x in ("договор", "контракт", "contract", "agreement", "соглашение")):
        return "contract_review"
    if any(x in lower for x in ("таблица", "xlsx", "csv", "excel", "выруч", "продаж", "аналитик", "аномал")):
        return "spreadsheet_analysis"
    if any(x in lower for x in ("кратко", "summary", "резюме", "сократи", "главное")):
        return "summary"
    if any(x in lower for x in ("ошибк", "проверь", "найди", "риск", "проблем")):
        return "risk_check"
    if any(x in lower for x in ("перепиши", "улучши", "сделай текст", "коммерческое", "кп", "письмо")):
        return "rewrite_or_business"
    return "general_analysis"


def build_document_profile(filename: str, extracted_text: str) -> str:
    ext = _extension(filename).upper() or "FILE"
    file_kind = detect_file_kind(filename, extracted_text)
    text = extracted_text or ""
    lines = [line for line in text.splitlines() if line.strip()]
    profile = [
        f"Формат: {ext}",
        f"Тип: {file_kind}",
        f"Объём извлечённого текста: {len(text)} символов",
        f"Непустых строк в выборке: {len(lines)}",
    ]
    if len(text) >= FILE_TEXT_LIMIT - 200:
        profile.append("Важно: файл был обрезан до лимита анализа, поэтому отвечай только по доступной части.")
    return "\n".join(f"• {item}" for item in profile)


def build_file_analysis_prompt(question: str, filename: str, extracted_text: str) -> str:
    file_kind = detect_file_kind(filename, extracted_text)
    task = detect_file_task(question, filename, extracted_text)
    question = question.strip() or "Проанализируй файл и выдели главное."

    task_rules = {
        "contract_review": (
            "Фокус: договор/юридический документ. Выдели стороны, предмет, деньги, сроки, обязанности, штрафы, "
            "расторжение, спорные места и риски. Не давай юридическую гарантию; формулируй как аналитический разбор."
        ),
        "spreadsheet_analysis": (
            "Фокус: таблица/данные. Найди ключевые показатели, заметные изменения, аномалии, возможные ошибки, "
            "выводы для бизнеса. Если точных расчётов нет в выборке — не придумывай цифры."
        ),
        "summary": "Фокус: краткое содержание. Дай сжатое резюме, главные мысли и что важно запомнить.",
        "risk_check": "Фокус: проверка. Найди ошибки, противоречия, риски, слабые места и что стоит уточнить.",
        "rewrite_or_business": "Фокус: деловая переработка. Помоги улучшить документ, сделать его яснее, сильнее и профессиональнее.",
        "general_analysis": "Фокус: общий анализ. Дай понятный разбор, основные выводы и полезные следующие шаги.",
    }

    return (
        "Ты — File AI PRO аналитик VSotah AI. Работай как сильный ассистент по документам, таблицам и файлам. "
        "Отвечай на русском, уверенно, полезно и без воды.\n\n"
        f"Имя файла: {filename}\n"
        f"Тип файла: {file_kind}\n"
        f"Тип задачи: {task}\n"
        f"Вопрос пользователя: {question}\n\n"
        "Профиль файла:\n"
        f"{build_document_profile(filename, extracted_text)}\n\n"
        "Правила качества:\n"
        "1. Не выдумывай факты, цифры, пункты договора или строки таблицы, которых нет в извлечённом содержимом.\n"
        "2. Сначала отвечай прямо по запросу пользователя, а не пересказывай весь файл подряд.\n"
        "3. Если файл большой и часть текста обрезана, честно укажи, что анализ сделан по доступной части.\n"
        "4. Для таблиц обращай внимание на заголовки, суммы, даты, пропуски, повторы и подозрительные значения.\n"
        "5. Для договоров выделяй риски, сроки, деньги, обязанности, штрафы и неясные формулировки.\n"
        "6. Для учебных/исследовательских файлов объясняй простым языком и выделяй структуру.\n"
        "7. Не пиши технические детали про API, лимиты, токены или внутреннюю обработку.\n"
        "8. Если данных не хватает, скажи, что именно нужно прислать или уточнить.\n\n"
        f"Специальная инструкция по задаче: {task_rules.get(task, task_rules['general_analysis'])}\n\n"
        "Формат ответа:\n"
        "• Короткий вывод\n"
        "• Главное по файлу\n"
        "• Риски/ошибки/важные места, если есть\n"
        "• Что можно сделать дальше\n\n"
        "Содержимое файла для анализа:\n"
        f"{_limit_text(extracted_text)}"
    )

def build_file_status_text(filename: str, stage: str) -> str:
    ext = _extension(filename).upper() or "FILE"

    if stage == "reading":
        return (
            "📄 Обрабатываю файл...\n"
            f"⚡ Читаю {ext}: {filename}"
        )

    if stage == "analyzing":
        return (
            "🧠 VSotah AI анализирует файл...\n"
            "📊 Извлекаю ключевые данные...\n"
            "✨ Формирую AI-выводы..."
        )

    if stage == "scanned_pdf":
        return (
            "📄 PDF похож на скан.\n"
            "🖼 Анализирую страницы как изображения..."
        )

    return "📎 Обрабатываю файл..."


def render_pdf_pages_to_images(file_bytes: bytes, max_pages: int = PDF_VISION_MAX_PAGES) -> list[bytes]:
    try:
        import fitz
    except Exception:
        return []

    images: list[bytes] = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_index in range(min(len(doc), max_pages)):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            images.append(pix.tobytes("jpeg"))
        doc.close()
    except Exception as e:
        print(f"PDF RENDER ERROR: {e}")
        return []

    return images


async def analyze_pdf_images_with_openai(
    question: str,
    filename: str,
    file_bytes: bytes,
    history: list[dict],
    selected_model: str = "gpt",
) -> str:
    """Vision fallback for scanned PDFs and image-only documents."""
    page_images = await asyncio.to_thread(render_pdf_pages_to_images, file_bytes, PDF_VISION_MAX_PAGES)
    if not page_images or not openai_client:
        return ""

    question = question.strip() or "Прочитай PDF как документ: извлеки видимый текст, сделай краткое резюме, найди важные пункты, риски и следующие шаги."
    content: list[dict[str, Any]] = [
        {"type": "text", "text": build_pdf_vision_prompt(filename, question)}
    ]

    for img in page_images:
        image_b64 = base64.b64encode(img).decode("utf-8")
        content.append(
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{image_b64}",
                    "detail": "low",
                },
            }
        )

    system_text = system_prompt(get_personality(selected_model))
    openai_messages: list[dict[str, Any]] = [{"role": "system", "content": system_text}]
    for msg in history[-VISION_HISTORY_LIMIT:]:
        if msg.get("role") in {"user", "assistant"} and msg.get("content"):
            openai_messages.append({"role": msg["role"], "content": msg["content"]})
    openai_messages.append({"role": "user", "content": content})

    try:
        response = await asyncio.wait_for(
            openai_client.chat.completions.create(
                model=OPENAI_VISION_MODEL,
                messages=openai_messages,
                temperature=0.4,
                max_completion_tokens=VISION_MAX_TOKENS,
            ),
            timeout=AI_TIMEOUT_SECONDS,
        )
        return clean_ai_answer(response.choices[0].message.content or "")
    except Exception as e:
        print(f"PDF VISION ERROR: {str(e).replace(chr(10), ' ')[:1200]}")
        return ""


async def file_router(
    selected_model: str,
    question: str,
    filename: str,
    extracted_text: str,
    history: list[dict],
) -> str:
    """Routes extracted file content through the user's selected text AI model."""
    prompt = build_file_analysis_prompt(question, filename, extracted_text)
    messages = [*history[-TEXT_HISTORY_LIMIT:], {"role": "user", "content": prompt}]
    return await ai_router(selected_model, messages)




